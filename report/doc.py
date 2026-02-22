"""
Copyright (c) 2025 Ed Millard

Permission is hereby granted, free of charge, to any person obtaining a copy
of this software and associated documentation files (the "Software"), to deal
in the Software without restriction, including without limitation the rights
to use, copy, modify, merge, publish, distribute copies of the Software, and
to permit persons to whom the Software is furnished to do so, subject to the
following conditions:

The above copyright notice and this permission notice shall be included in all
copies or substantial portions of the Software.

THE SOFTWARE IS PROVIDED "AS IS", WITHOUT WARRANTY OF ANY KIND, EXPRESS OR
IMPLIED, INCLUDING BUT NOT LIMITED TO THE WARRANTIES OF MERCHANTABILITY,
FITNESS FOR A PARTICULAR PURPOSE AND NONINFRINGEMENT. IN NO EVENT SHALL THE
AUTHORS OR COPYRIGHT HOLDERS BE LIABLE FOR ANY CLAIM, DAMAGES OR OTHER
LIABILITY, WHETHER IN AN ACTION OF CONTRACT, TORT OR OTHERWISE, ARISING FROM,
OUT OF OR IN CONNECTION WITH THE SOFTWARE OR THE USE OR OTHER DEALINGS IN THE
SOFTWARE.
"""
import os
from pathlib import Path
import platform
import subprocess
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import re

class Report:
    def __init__(self):
        self.doc = Document()

    def save_doc(self, file_name):
        # 7. SAVE DOCUMENT
        self.doc.save(file_name)
        print(f'Report saved as {file_name}')

    def page_setup(self):
        section = self.doc.sections[0]  # Default section

        # section.orientation = WD_ORIENT.PORTRAIT
        section.orientation = WD_ORIENT.LANDSCAPE
        # Swap page width and height for landscape (US Letter)
        section.page_width = Inches(11)
        section.page_height = Inches(8.5)
        # Set margins (1 inch each side)
        section.left_margin = section.right_margin = Inches(1)
        section.top_margin = section.bottom_margin = Inches(1)

    def create_title_page(self, year, author):
        # Set default font
        style = self.doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(12)

        for _ in range(5):
            self.doc.add_paragraph()

        # Add title (three lines)
        title_lines = ['Dolores Project', 'Water Accounting Report', str(year)]

        for i, line in enumerate(title_lines):
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line)
            run.bold = True
            # Set font size: larger for first two lines, slightly smaller for year
            run.font.size = Pt(24) if i < 2 else Pt(20)
            # Add spacing after each line
            p.space_after = Inches(0.2)

        # Add two blank lines
        for _ in range(2):
            self.doc.add_paragraph()

        # Add author
        author_p = self.doc.add_paragraph()
        author_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        author_run = author_p.add_run(f'Prepared by: {author}')
        author_run.font.size = Pt(14)
        author_p.space_after = Inches(0.2)

        # Add current date
        current_date = datetime.now().strftime('%B %d, %Y')
        date_p = self.doc.add_paragraph()
        date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_p.add_run(f'Date: {current_date}')
        date_run.font.size = Pt(14)

        # Add some bottom margin
        self.doc.add_paragraph()

    def header(self, header, header_level=1):
        self.doc.add_heading(header, level=header_level)

    def paragraph_with_header(self, header, text, header_level=3, font=None, size=None):
        self.doc.add_heading(header, level=header_level)
        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run(text)

        if font is not None:
            run.font.name = 'Courier New'  # Cross-platform monospaced font

        if size is not None:
            run.font.size = Pt(size)

    def page_break(self):
        self.doc.add_page_break()

    def paragraph(self, text, font=None, size=None):
        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run(text)
        if font is not None:
            run.font.name = 'Courier New'  # Cross-platform monospaced font

        if size is not None:
            run.font.size = Pt(size)

    def plot(self, plot_image_file_name, inches=6):
        self.doc.add_picture(plot_image_file_name, width=Inches(inches))

        # header = self.doc.sections[0].header
        # header_para = header.paragraphs[0]
        # header_para.text = "Water Flow Analysis Report"
        # header_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        # header_para.runs[0].bold = True
        # header_para.runs[0].font.size = Inches(0.18)

        # Remove plot image
        os.remove(plot_image_file_name)

    @staticmethod
    def open_docx_in_app(file_path: Path):
        """
        Open a .docx file in LibreOffice Writer (Linux) or Microsoft Word (Windows).

        Args:
            file_path (Path): Path to the .docx file.

        Raises:
            FileNotFoundError: If the file doesn't exist.
            OSError: If the application can't be launched.
        """
        file_path = file_path.resolve()  # Resolve to absolute path
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        system = platform.system().lower()

        if system == "linux":
            # Launch LibreOffice Writer with the file
            try:
                subprocess.run(["libreoffice", "--writer", str(file_path)], check=True)
                print(f"Opened {file_path} in LibreOffice Writer.")
            except subprocess.CalledProcessError as e:
                raise OSError(f"Failed to launch LibreOffice: {e}")
            except FileNotFoundError:
                raise OSError("LibreOffice not found. Install it with 'sudo apt install libreoffice' (or equivalent).")

        elif system == "windows":
            # Open the file in the default app (Microsoft Word)
            try:
                os.startfile(str(file_path))
                print(f"Opened {file_path} in Microsoft Word or Excel.")
            except OSError as e:
                raise OSError(
                    f"Failed to launch Word: {e}. Ensure Microsoft Word is installed and associated with .docx files.")

        else:
            raise OSError(f"Unsupported OS: {system}. Only Linux and Windows are supported.")

    from docx import Document
    import re

    def add_markdown_paragraph(self, text, tag_pattern=r'\*\*(.*?)\*\*', font=None, size=None):
        """
        Adds a paragraph to the document, parsing the input text for bold tags
        (default: **text**) and applying bold formatting to matched substrings.

        Args:
            doc: The Document object to add to.
            text (str): The input string with optional **bold** tags.
            tag_pattern (str): Regex pattern for tags (default matches **bold**).
            font (str): Font for run
            size (int): Point size for run

        Returns:
            Paragraph: The created paragraph with formatted runs.
        """
        paragraph = self.doc.add_paragraph()

        # Split text into parts: non-bold (even indices), bold (odd indices)
        parts = re.split(tag_pattern, text)
        for i, part in enumerate(parts):
            if not part:
                continue
            if i % 2 == 1:  # Odd indices are bold content (captured group)
                run = paragraph.add_run(part)
                if font is not None:
                    run.font.name = font
                if size is not None:
                    run.font.size = Pt(size)
                run.bold = True
            else:  # Even indices are plain text
                run = paragraph.add_run(part)
                if font is not None:
                    run.font.name = font
                if size is not None:
                    run.font.size = Pt(size)

        return paragraph

    def table_of_contents(self):
        # This may only work in Windows Word and is untested,
        # In LibreOffice currently you create the TOC manually using inser
        #
        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run()

        # Create TOC field elements
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'begin')

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        # TOC instruction: \o "2-2" for only level 2 headings; \h for hyperlinks; \z for table entry field; \u for user TOC
        instrText.text = 'TOC \\o "2-2" \\h \\z \\u'

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')

        fldChar3 = OxmlElement('w:t')
        fldChar3.text = "Right-click to update field."

        fldChar2.append(fldChar3)

        fldChar4 = OxmlElement('w:fldChar')
        fldChar4.set(qn('w:fldCharType'), 'end')

        # Append elements to the run
        r_element = run._r
        r_element.append(fldChar)
        r_element.append(instrText)
        r_element.append(fldChar2)
        r_element.append(fldChar4)